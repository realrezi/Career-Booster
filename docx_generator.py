import io
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def generate_cv_docx(cv_data: dict, theme: str = "modern") -> io.BytesIO:
    """Generates an editable MS Word (.docx) document formatted matching the requested theme."""
    doc = docx.Document()
    
    # Page Margins (0.6 inch / ~15mm)
    sections = doc.sections
    for sec in sections:
        sec.top_margin = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.6)
        sec.right_margin = Inches(0.6)

    # Theme Fonts & Colors
    font_name = "Times New Roman" if theme == "academic" else "Arial"
    primary_color = RGBColor(15, 23, 42)    # Dark slate navy
    secondary_color = RGBColor(71, 85, 105) # Slate gray
    body_color = RGBColor(30, 41, 59)      # Charcoal

    # Base Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = font_name
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = body_color

    # --- HEADER ---
    name_str = cv_data.get("name", "Candidate Name").strip()
    if name_str.endswith("|"): name_str = name_str[:-1].strip()
    
    p_name = doc.add_paragraph()
    if theme == "academic":
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    r_name = p_name.add_run(name_str)
    r_name.font.name = font_name
    r_name.font.size = Pt(18)
    r_name.font.bold = True
    r_name.font.color.rgb = primary_color
    p_name.paragraph_format.space_after = Pt(2)

    # Subtitle / Title
    title_str = cv_data.get("title")
    if title_str:
        if title_str.endswith("|"): title_str = title_str[:-1].strip()
        p_title = doc.add_paragraph()
        if theme == "academic":
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_title = p_title.add_run(title_str)
        r_title.font.name = font_name
        r_title.font.size = Pt(10.5)
        r_title.font.bold = (theme == "modern")
        r_title.font.italic = (theme == "academic")
        r_title.font.color.rgb = secondary_color
        p_title.paragraph_format.space_after = Pt(2)

    # Contact Info
    contact = cv_data.get('contact', {})
    contact_parts = []
    if contact.get('location'): contact_parts.append(contact.get('location').strip())
    if contact.get('email'): contact_parts.append(contact.get('email').strip())
    if contact.get('phone'): contact_parts.append(contact.get('phone').strip())

    links = contact.get('links', [])
    if isinstance(links, list):
        for l in links:
            if isinstance(l, str) and l.strip(): contact_parts.append(l.strip())
    elif isinstance(links, str) and links.strip():
        contact_parts.append(links.strip())

    clean_parts = [p.strip(" |") for p in contact_parts if p and p.strip(" |")]
    if clean_parts:
        p_contact = doc.add_paragraph()
        if theme == "academic":
            p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_contact = p_contact.add_run("  |  ".join(clean_parts))
        r_contact.font.name = font_name
        r_contact.font.size = Pt(9)
        r_contact.font.color.rgb = secondary_color
        p_contact.paragraph_format.space_after = Pt(10)

    def add_section_header(title_text):
        p_sec = doc.add_paragraph()
        p_sec.paragraph_format.space_before = Pt(8)
        p_sec.paragraph_format.space_after = Pt(4)
        if theme == "academic":
            p_sec.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        r_sec = p_sec.add_run(title_text.upper())
        r_sec.font.name = font_name
        r_sec.font.size = Pt(11)
        r_sec.font.bold = True
        r_sec.font.color.rgb = primary_color

    # --- PROFESSIONAL PROFILE ---
    profile = cv_data.get('professional_profile')
    if profile and profile.strip():
        add_section_header("Professional Profile")
        p_prof = doc.add_paragraph(profile.strip())
        p_prof.paragraph_format.space_after = Pt(6)

    # --- DYNAMIC SECTIONS ---
    for section in cv_data.get('sections', []):
        sec_title = section.get('section_title', '').strip()
        items = section.get('items', [])
        
        if sec_title and items:
            add_section_header(sec_title)
            
            for item in items:
                if isinstance(item, str):
                    p_b = doc.add_paragraph(style='List Bullet')
                    p_b.paragraph_format.space_after = Pt(2)
                    r_b = p_b.add_run(item.strip())
                    r_b.font.name = font_name
                elif isinstance(item, dict):
                    title_date = item.get('title_date', '').strip()
                    if title_date:
                        has_bullets = bool(item.get('bullets'))
                        is_short_title = len(title_date) < 80
                        
                        p_t = doc.add_paragraph()
                        p_t.paragraph_format.space_after = Pt(2)
                        r_t = p_t.add_run(title_date)
                        r_t.font.name = font_name
                        if has_bullets or is_short_title:
                            r_t.font.bold = True
                            r_t.font.color.rgb = primary_color
                    
                    for bullet in item.get('bullets', []):
                        if bullet and bullet.strip():
                            p_b = doc.add_paragraph(style='List Bullet')
                            p_b.paragraph_format.space_after = Pt(2)
                            r_b = p_b.add_run(bullet.strip())
                            r_b.font.name = font_name

    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream
